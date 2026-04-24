import io
import logging
import re
from datetime import datetime, UTC
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

from backend.prompts import extract_chapter_title


logger = logging.getLogger(__name__)


def _approved_chapters(chapters: list[dict]) -> list[dict]:
    return sorted(
        [chapter for chapter in chapters if chapter.get("status") == "approved"],
        key=lambda chapter: chapter["chapter_number"],
    )


def _chapter_entries(outline: dict, chapters: list[dict]) -> list[tuple[int, str]]:
    return [
        (
            chapter["chapter_number"],
            extract_chapter_title(outline.get("content", ""), chapter["chapter_number"]),
        )
        for chapter in _approved_chapters(chapters)
    ]


def _overview_lines(book: dict, approved_count: int, generated_on: str) -> list[tuple[str, str]]:
    return [
        ("Title", book["title"]),
        ("Prepared", generated_on),
        ("Status", book.get("status", "unknown")),
        ("Approved Chapters", str(approved_count)),
    ]


def compile_to_docx(book: dict, outline: dict, chapters: list) -> bytes:
    """
    Assemble a formatted .docx manuscript in memory and return the raw bytes.

    The document contains a title page, author notes, outline, and all approved
    chapters in chapter order with optional summary blocks.
    """
    try:
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        normal_style = document.styles["Normal"]
        normal_style.font.name = "Times New Roman"
        normal_style.font.size = Pt(12)

        generated_on = datetime.now(UTC).strftime("%B %d, %Y")
        author_name = book.get("author") or "AutoBook"
        approved_chapters = _approved_chapters(chapters)
        chapter_entries = _chapter_entries(outline, chapters)

        title_paragraph = document.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(book["title"])
        title_run.bold = True
        title_run.font.size = Pt(28)

        subtitle_paragraph = document.add_paragraph()
        subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_paragraph.add_run(f"By {author_name}")
        subtitle_run.italic = True
        subtitle_run.font.size = Pt(14)

        date_paragraph = document.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_paragraph.add_run(f"Generated on {generated_on}").font.size = Pt(11)

        cover_footer = section.footer.paragraphs[0]
        cover_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cover_footer.text = book["title"]

        document.add_page_break()

        toc_heading = document.add_paragraph()
        toc_heading.style = document.styles["Heading 1"]
        toc_heading.add_run("Table of Contents")

        for chapter_number, chapter_title in chapter_entries:
            toc_entry = document.add_paragraph()
            toc_entry.paragraph_format.left_indent = Inches(0.25)
            toc_entry.paragraph_format.space_after = Pt(6)
            toc_entry.add_run(f"Chapter {chapter_number} - {chapter_title}")

        document.add_page_break()

        overview_heading = document.add_paragraph()
        overview_heading.style = document.styles["Heading 1"]
        overview_heading.add_run("Manuscript Overview")
        for label, value in _overview_lines(book, len(approved_chapters), generated_on):
            overview_paragraph = document.add_paragraph()
            overview_paragraph.paragraph_format.space_after = Pt(8)
            overview_paragraph.add_run(f"{label}: ").bold = True
            overview_paragraph.add_run(value)

        document.add_page_break()

        notes_heading = document.add_paragraph()
        notes_heading.style = document.styles["Heading 1"]
        notes_heading.add_run("Author Notes")
        for paragraph in re.split(r"\n\s*\n", (book.get("notes") or "").strip()):
            if paragraph.strip():
                note_paragraph = document.add_paragraph(paragraph.strip())
                note_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                note_paragraph.paragraph_format.first_line_indent = Inches(0.3)
                note_paragraph.paragraph_format.space_after = Pt(10)

        document.add_page_break()

        outline_heading = document.add_paragraph()
        outline_heading.style = document.styles["Heading 1"]
        outline_heading.add_run("Outline")
        for paragraph in re.split(r"\n\s*\n", (outline.get("content") or "").strip()):
            if paragraph.strip():
                outline_paragraph = document.add_paragraph(paragraph.strip())
                outline_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                outline_paragraph.paragraph_format.first_line_indent = Inches(0.3)
                outline_paragraph.paragraph_format.space_after = Pt(10)

        for index, chapter in enumerate(approved_chapters):
            document.add_section(WD_SECTION_START.NEW_PAGE)
            current_section = document.sections[-1]
            current_section.top_margin = Inches(1)
            current_section.bottom_margin = Inches(1)
            current_section.left_margin = Inches(1)
            current_section.right_margin = Inches(1)

            header = current_section.header.paragraphs[0]
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            header.text = book["title"]

            footer = current_section.footer.paragraphs[0]
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer.text = f"Chapter {chapter['chapter_number']}"

            chapter_heading = document.add_paragraph()
            chapter_heading.style = document.styles["Heading 1"]
            chapter_heading.add_run(
                f"Chapter {chapter_entries[index][0]} - {chapter_entries[index][1]}"
            )

            for paragraph in re.split(r"\n\s*\n", (chapter.get("content") or "").strip()):
                if paragraph.strip():
                    body_paragraph = document.add_paragraph(paragraph.strip())
                    body_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    body_paragraph.paragraph_format.first_line_indent = Inches(0.3)
                    body_paragraph.paragraph_format.line_spacing = 1.3
                    body_paragraph.paragraph_format.space_after = Pt(10)

            if chapter.get("summary"):
                summary_paragraph = document.add_paragraph()
                summary_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                summary_paragraph.paragraph_format.space_before = Pt(8)
                summary_paragraph.paragraph_format.space_after = Pt(8)
                summary_run = summary_paragraph.add_run(
                    f"Summary: {chapter['summary'].strip()}"
                )
                summary_run.italic = True

        output = io.BytesIO()
        document.save(output)
        return output.getvalue()
    except Exception as exc:
        logger.error("DOCX compilation failed: %s", exc, exc_info=True)
        raise


def compile_to_pdf(book: dict, outline: dict, chapters: list) -> bytes:
    """
    Assemble a formatted PDF manuscript in memory and return the raw bytes.

    The document contains a title page, author notes, outline, and all approved
    chapters in chapter order with optional summary blocks.
    """
    try:
        output = io.BytesIO()

        generated_on = datetime.now(UTC).strftime("%B %d, %Y")
        author_name = book.get("author") or "AutoBook"
        approved_chapters = _approved_chapters(chapters)
        chapter_entries = _chapter_entries(outline, chapters)

        document = SimpleDocTemplate(
            output,
            pagesize=LETTER,
            leftMargin=inch,
            rightMargin=inch,
            topMargin=inch,
            bottomMargin=inch,
        )
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "BookTitle",
            parent=styles["Title"],
            fontSize=28,
            leading=34,
            alignment=1,
            spaceAfter=16,
            fontName="Times-Bold",
        )
        subtitle_style = ParagraphStyle(
            "BookSubtitle",
            parent=styles["BodyText"],
            fontSize=14,
            leading=18,
            alignment=1,
            italic=True,
            spaceAfter=8,
        )
        heading_style = ParagraphStyle(
            "SectionHeading",
            parent=styles["Heading1"],
            fontSize=18,
            leading=22,
            spaceAfter=12,
            fontName="Times-Bold",
        )
        body_style = ParagraphStyle(
            "BodyCopy",
            parent=styles["BodyText"],
            fontSize=12,
            leading=17,
            spaceAfter=10,
            fontName="Times-Roman",
            alignment=4,
            firstLineIndent=18,
        )
        summary_style = ParagraphStyle(
            "SummaryCopy",
            parent=styles["BodyText"],
            fontSize=10,
            leading=14,
            spaceAfter=10,
            italic=True,
            leftIndent=18,
            rightIndent=18,
            fontName="Times-Italic",
        )
        toc_style = ParagraphStyle(
            "TOCEntry",
            parent=styles["BodyText"],
            fontSize=12,
            leading=16,
            spaceAfter=8,
            leftIndent=18,
            fontName="Times-Roman",
        )
        overview_style = ParagraphStyle(
            "OverviewCopy",
            parent=styles["BodyText"],
            fontSize=12,
            leading=16,
            spaceAfter=8,
            fontName="Times-Bold",
        )

        story = [
            Spacer(1, 2.2 * inch),
            Paragraph(escape(book["title"]), title_style),
            Paragraph(escape(f"By {author_name}"), subtitle_style),
            Paragraph(escape(f"Generated on {generated_on}"), styles["BodyText"]),
            PageBreak(),
            Paragraph("Table of Contents", heading_style),
        ]

        for chapter_number, chapter_title in chapter_entries:
            story.append(
                Paragraph(
                    escape(f"Chapter {chapter_number} - {chapter_title}"),
                    toc_style,
                )
            )

        story.extend([PageBreak(), Paragraph("Manuscript Overview", heading_style)])
        for label, value in _overview_lines(book, len(approved_chapters), generated_on):
            story.append(
                Paragraph(
                    f"{escape(label)}: <font name='Times-Roman'>{escape(value)}</font>",
                    overview_style,
                )
            )

        story.extend([PageBreak(), Paragraph("Author Notes", heading_style)])

        for paragraph in re.split(r"\n\s*\n", (book.get("notes") or "").strip()):
            if paragraph.strip():
                story.append(
                    Paragraph(
                        escape(paragraph.strip()).replace("\n", "<br/>"),
                        body_style,
                    )
                )

        story.extend([PageBreak(), Paragraph("Outline", heading_style)])

        for paragraph in re.split(r"\n\s*\n", (outline.get("content") or "").strip()):
            if paragraph.strip():
                story.append(
                    Paragraph(
                        escape(paragraph.strip()).replace("\n", "<br/>"),
                        body_style,
                    )
                )

        def draw_page_chrome(canvas, doc):
            if canvas.getPageNumber() > 1:
                canvas.saveState()
                canvas.setFont("Times-Roman", 9)
                canvas.drawCentredString(LETTER[0] / 2.0, 0.6 * inch, book["title"])
                canvas.drawRightString(LETTER[0] - inch, 0.6 * inch, str(canvas.getPageNumber()))
                canvas.restoreState()

        for index, chapter in enumerate(approved_chapters):
            story.append(PageBreak())
            story.append(
                Paragraph(
                    escape(
                        f"Chapter {chapter_entries[index][0]} - {chapter_entries[index][1]}"
                    ),
                    heading_style,
                )
            )

            for paragraph in re.split(r"\n\s*\n", (chapter.get("content") or "").strip()):
                if paragraph.strip():
                    story.append(
                        Paragraph(
                            escape(paragraph.strip()).replace("\n", "<br/>"),
                            body_style,
                        )
                    )

            if chapter.get("summary"):
                story.append(Paragraph(escape(chapter["summary"].strip()), summary_style))

        document.build(
            story,
            onFirstPage=draw_page_chrome,
            onLaterPages=draw_page_chrome,
        )
        return output.getvalue()
    except Exception as exc:
        logger.error("PDF compilation failed: %s", exc, exc_info=True)
        raise


def compile_to_txt(book: dict, outline: dict, chapters: list) -> bytes:
    generated_on = datetime.now(UTC).strftime("%B %d, %Y")
    approved_chapters = _approved_chapters(chapters)

    compiled_sections = [
        book["title"].strip(),
        "=" * max(len(book["title"].strip()), 12),
        "",
        "MANUSCRIPT OVERVIEW",
        "-" * 19,
    ]

    for label, value in _overview_lines(book, len(approved_chapters), generated_on):
        compiled_sections.append(f"{label}: {value}")

    if book.get("notes"):
        compiled_sections.extend(["", "AUTHOR NOTES", "-" * 12, book["notes"].strip()])

    if outline.get("content"):
        compiled_sections.extend(["", "OUTLINE", "-" * 7, outline["content"].strip()])

    for chapter in approved_chapters:
        chapter_title = extract_chapter_title(
            outline.get("content", ""), chapter["chapter_number"]
        )
        heading = f"CHAPTER {chapter['chapter_number']}: {chapter_title}"
        compiled_sections.extend(["", heading, "-" * len(heading), chapter["content"].strip()])
        if chapter.get("summary"):
            compiled_sections.extend(["", f"Chapter Summary: {chapter['summary'].strip()}"])

    return "\n".join(compiled_sections).strip().encode("utf-8")
