import io
import logging
import re
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

from prompts import extract_chapter_title


logger = logging.getLogger(__name__)


def compile_to_docx(book: dict, outline: dict, chapters: list) -> bytes:
    """
    Assemble a formatted .docx manuscript in memory and return the raw bytes.

    The document contains a title page, author notes, outline, and all approved
    chapters in chapter order with optional summary blocks.
    """
    try:
        document = Document()

        title_paragraph = document.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(book["title"])
        title_run.bold = True
        title_run.font.size = Pt(28)

        document.add_page_break()

        document.add_heading("Author Notes", level=1)
        for paragraph in re.split(r"\n\s*\n", (book.get("notes") or "").strip()):
            if paragraph.strip():
                document.add_paragraph(paragraph.strip())

        document.add_page_break()

        document.add_heading("Outline", level=1)
        for paragraph in re.split(r"\n\s*\n", (outline.get("content") or "").strip()):
            if paragraph.strip():
                document.add_paragraph(paragraph.strip())

        document.add_page_break()

        approved_chapters = sorted(
            [chapter for chapter in chapters if chapter.get("status") == "approved"],
            key=lambda chapter: chapter["chapter_number"],
        )

        for index, chapter in enumerate(approved_chapters):
            chapter_title = extract_chapter_title(
                outline.get("content", ""), chapter["chapter_number"]
            )
            document.add_heading(
                f"Chapter {chapter['chapter_number']} - {chapter_title}",
                level=1,
            )

            for paragraph in re.split(r"\n\s*\n", (chapter.get("content") or "").strip()):
                if paragraph.strip():
                    document.add_paragraph(paragraph.strip())

            if chapter.get("summary"):
                summary_paragraph = document.add_paragraph()
                summary_run = summary_paragraph.add_run(chapter["summary"].strip())
                summary_run.italic = True

            if index < len(approved_chapters) - 1:
                document.add_page_break()

        output = io.BytesIO()
        document.save(output)
        return output.getvalue()
    except Exception as exc:
        logger.error("DOCX compilation failed: %s", exc, exc_info=True)
        raise


def compile_to_pdf(book: dict, outline: dict, chapters: list) -> bytes:
    """
    Assemble a formatted PDF manuscript in memory and return the raw bytes.

    The document contains a title page, author notes, outline, and all approved
    chapters in chapter order with optional summary blocks.
    """
    try:
        output = io.BytesIO()
        document = SimpleDocTemplate(
            output,
            pagesize=LETTER,
            leftMargin=inch,
            rightMargin=inch,
            topMargin=inch,
            bottomMargin=inch,
        )
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "BookTitle",
            parent=styles["Title"],
            fontSize=28,
            leading=34,
            alignment=1,
            spaceAfter=0,
        )
        heading_style = ParagraphStyle(
            "SectionHeading",
            parent=styles["Heading1"],
            fontSize=18,
            leading=22,
            spaceAfter=12,
        )
        body_style = ParagraphStyle(
            "BodyCopy",
            parent=styles["BodyText"],
            fontSize=12,
            leading=16,
            spaceAfter=10,
        )
        summary_style = ParagraphStyle(
            "SummaryCopy",
            parent=styles["BodyText"],
            fontSize=10,
            leading=14,
            spaceAfter=10,
            italic=True,
        )

        story = [
            Spacer(1, 2.5 * inch),
            Paragraph(escape(book["title"]), title_style),
            PageBreak(),
            Paragraph("Author Notes", heading_style),
        ]

        for paragraph in re.split(r"\n\s*\n", (book.get("notes") or "").strip()):
            if paragraph.strip():
                story.append(
                    Paragraph(
                        escape(paragraph.strip()).replace("\n", "<br/>"),
                        body_style,
                    )
                )

        story.extend([PageBreak(), Paragraph("Outline", heading_style)])

        for paragraph in re.split(r"\n\s*\n", (outline.get("content") or "").strip()):
            if paragraph.strip():
                story.append(
                    Paragraph(
                        escape(paragraph.strip()).replace("\n", "<br/>"),
                        body_style,
                    )
                )

        story.append(PageBreak())

        approved_chapters = sorted(
            [chapter for chapter in chapters if chapter.get("status") == "approved"],
            key=lambda chapter: chapter["chapter_number"],
        )

        for index, chapter in enumerate(approved_chapters):
            chapter_title = extract_chapter_title(
                outline.get("content", ""), chapter["chapter_number"]
            )
            story.append(
                Paragraph(
                    escape(f"Chapter {chapter['chapter_number']} - {chapter_title}"),
                    heading_style,
                )
            )

            for paragraph in re.split(r"\n\s*\n", (chapter.get("content") or "").strip()):
                if paragraph.strip():
                    story.append(
                        Paragraph(
                            escape(paragraph.strip()).replace("\n", "<br/>"),
                            body_style,
                        )
                    )

            if chapter.get("summary"):
                story.append(Paragraph(escape(chapter["summary"].strip()), summary_style))

            if index < len(approved_chapters) - 1:
                story.append(PageBreak())

        document.build(story)
        return output.getvalue()
    except Exception as exc:
        logger.error("PDF compilation failed: %s", exc, exc_info=True)
        raise
